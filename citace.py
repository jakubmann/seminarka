from isbnlib import *
import isbnlib
from docx import *

query = input('Zadejte text obsažený v knize (isbn, jméno autora, citát): ')
place_published = input('Zadejte míto vydání(není v databázi): ')


isbn = isbn_from_words(query)


book = meta(isbn)

if not book:
    exit()

def formatAuthor():
    for author in book['Authors']:
        author = author.split()
        return ('%s, %s' % (author[1].upper(), author[0]))

def formatISBN():
    return mask(isbn, separator='-')

document = Document()

p = document.add_paragraph('%s, %s. ' % (formatAuthor(), book['Year']))
p.add_run(book['Title']).italic = True
p.add_run('. %s: %s. ISBN %s' % (place_published, book['Publisher'], formatISBN()))

document.save('citace.docx')
